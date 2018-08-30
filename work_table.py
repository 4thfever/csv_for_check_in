import pandas as pd 
from find_workday import DaysInfo
from docx import Document

def output_docx(year,name,**args):
	c = DaysInfo(year,**args)
	df_days = c.df_days
	# print(c.df_days)
	columns = [i for i in range(32)]
	df_days_only = df_days[df_days.columns[0:31]]
	df_info_only = df_days[df_days.columns[31:]]
	df_days_only = df_days_only.replace(0,'/')
	df_days_only = df_days_only.replace(1,chr(ord('√')))
	df_days_only = df_days_only.replace(-1,'')
	df_days_only = df_days_only.replace(3,chr(ord('☆')))
	df_days_only = df_days_only.replace(4,'◇')
	df_days_only = df_days_only.replace(5,'※')
	df_days_only = df_days_only.replace(6,'□')
	df_days_only = df_days_only.replace(7,'○')
	df_days_only = df_days_only.replace(8,'Δ')
	df_days_only = df_days_only.replace(9,'▽')
	df_days_only = df_days_only.replace(10,'♀')
	df_days_only = df_days_only.replace(11,'⊥')
	df_days_only = df_days_only.replace(12,'＃')
	df_days_only = df_days_only.replace(13,'×')
	df_days = pd.concat([df_days_only,df_info_only],axis =1)

	# df_days.to_csv('output.csv',encoding ='gbk')

	doc = Document('example.docx')
	table = doc.tables[0]
	for i,row in enumerate(table.rows):
		for j,cell in enumerate(row.cells):
			# print(cell.text)
			if i>1:
				if '年' in cell.text:
					cell.text = cell.text.replace('年','{}年'.format(year))
				if j>1:
					month_num = i-2
					day_num = j-2
					#从第二行第二列起放一月一日的数据
					cell.text = str(df_days.values[month_num][day_num])
				if j == 1:
					cell.text = name


	old_string = doc.paragraphs[0].text
	inline = doc.paragraphs[0].runs
	for i in range(len(inline)):
		inline[i].text = old_string.replace('年','{}年'.format(year))
	
	old_string = doc.paragraphs[1].text
	inline = doc.paragraphs[1].runs
	for i in range(len(inline)):
		inline[i].text = old_string.replace('姓名： ','姓名： {}'.format(name))

	doc.save('output.docx')
if __name__ == '__main__':
	output_docx(2012,'王心桥')