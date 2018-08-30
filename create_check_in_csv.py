'''
读取日程信息，生成符合联研院打卡规范的docx考勤表
'''
import pandas as pd 
from find_workday import DaysInfo
from docx import Document

#输入年份，姓名和其他信息，输出docx考勤表
def create_check_in_csv(year,name,**args):
	#调用find_workday程序
	c = DaysInfo(year,**args)
	df_days = c.df_days
	
	#将表分为考勤部分和统计部分，分别处理
	columns = [i for i in range(32)]
	df_days_only = df_days[df_days.columns[0:31]]
	df_info_only = df_days[df_days.columns[31:]]

	#将数字换为图形
	df_days_only = df_days_only.replace(0,'/')
	df_days_only = df_days_only.replace(1,chr(ord('√')))
	df_days_only = df_days_only.replace(-1,'')
	df_days_only = df_days_only.replace(3,chr(ord('☆')))
	df_days_only = df_days_only.replace(4,'◇')
	df_days_only = df_days_only.replace(5,'※')
	df_days_only = df_days_only.replace(6,'□')
	df_days_only = df_days_only.replace(7,'○')
	df_days_only = df_days_only.replace(8,'Δ')
	df_days_only = df_days_only.replace(9,'▽')
	df_days_only = df_days_only.replace(10,'♀')
	df_days_only = df_days_only.replace(11,'⊥')
	df_days_only = df_days_only.replace(12,'＃')
	df_days_only = df_days_only.replace(13,'×')
	#重新合并
	df_days = pd.concat([df_days_only,df_info_only],axis =1)

	# 打开模板
	doc = Document('example.docx')
	#读入唯一的表格，逐行检索
	table = doc.tables[0]
	for i,row in enumerate(table.rows):
		for j,cell in enumerate(row.cells):
			# print(cell.text)
			if i>1:
				if '年' in cell.text:
					cell.text = cell.text.replace('年','{}年'.format(year))
				#从第二行第二列起放一月一日的数据
				if j>1:
					month_num = i-2
					day_num = j-2
					
					cell.text = str(df_days.values[month_num][day_num])
				#放姓名
				if j == 1:
					cell.text = name
	# 替换数据，为了维持格式不变，采用run的方式
	old_string = doc.paragraphs[0].text
	inline = doc.paragraphs[0].runs
	for i in range(len(inline)):
		inline[i].text = old_string.replace('年','{}年'.format(year))
	
	old_string = doc.paragraphs[1].text
	inline = doc.paragraphs[1].runs
	for i in range(len(inline)):
		inline[i].text = old_string.replace('姓名： ','姓名： {}'.format(name))

	doc.save('output.docx')
if __name__ == '__main__':
	# create_check_in_csv(2012,'王心桥')
	create_check_in_csv(2018,'王某某',chuchai = [[1,1],[12,31]],zhuwai = [[1,4]],kuanggong = [[1,2]])